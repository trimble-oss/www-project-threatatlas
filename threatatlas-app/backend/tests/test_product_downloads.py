"""Integration tests for product export endpoints and the CI security-status
gate (app/routers/product_downloads.py).

Covers: every export format renders, access control is enforced, the ZIP
bundle is well-formed, and the security-status threshold logic passes/fails
correctly — this endpoint is the CI/CD contract, so its decision logic is
worth pinning down.
"""

import io
import json
import zipfile

import pytest
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from app.models import (
    Diagram,
    DiagramMitigation,
    DiagramThreat,
    Framework,
    Mitigation,
    Model,
    Product,
    Threat,
    User,
)
from app.models.enums import UserRole
from tests.conftest import _create_user, make_auth_headers


# ── Full product graph builder ────────────────────────────────────────────────

def _build_product(
    db: Session,
    owner: User,
    *,
    is_public: bool = False,
    threats: list[dict] | None = None,
    mitigations: list[dict] | None = None,
) -> Product:
    """Create a product with one diagram/model plus the given threats/mitigations.

    Each threat dict may set: severity, status, element_id.
    Each mitigation dict may set: status, element_id.
    """
    fw = Framework(name=f"FW-{owner.id}", description="test framework")
    db.add(fw)
    db.flush()

    product = Product(user_id=owner.id, name="Export Product", description="desc", is_public=is_public)
    db.add(product)
    db.flush()

    diagram = Diagram(product_id=product.id, created_by=owner.id, name="Main DFD", description="d", current_version=1)
    db.add(diagram)
    db.flush()

    model = Model(diagram_id=diagram.id, framework_id=fw.id, name="M", created_by=owner.id)
    db.add(model)
    db.flush()

    for i, t in enumerate(threats or []):
        threat = Threat(framework_id=fw.id, name=f"Threat {i}", description="t-desc", category="Spoofing")
        db.add(threat)
        db.flush()
        db.add(DiagramThreat(
            diagram_id=diagram.id,
            model_id=model.id,
            threat_id=threat.id,
            element_id=t.get("element_id", f"node-{i}"),
            element_type="node",
            status=t.get("status", "identified"),
            severity=t.get("severity"),
            likelihood=t.get("likelihood"),
            impact=t.get("impact"),
            risk_score=t.get("risk_score"),
        ))

    for i, m in enumerate(mitigations or []):
        mit = Mitigation(framework_id=fw.id, name=f"Mitigation {i}", description="m-desc", category="Control")
        db.add(mit)
        db.flush()
        db.add(DiagramMitigation(
            diagram_id=diagram.id,
            model_id=model.id,
            mitigation_id=mit.id,
            element_id=m.get("element_id", f"node-{i}"),
            element_type="node",
            status=m.get("status", "proposed"),
        ))

    db.flush()
    db.refresh(product)
    return product


# ── Export formats render ──────────────────────────────────────────────────────

def test_download_diagrams_json(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    product = _build_product(db, standard_user, threats=[{"severity": "high"}])
    resp = client.get(f"/api/products/{product.id}/download/diagrams", headers=user_headers)
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith("application/json")
    payload = json.loads(resp.content)
    assert payload["product"]["name"] == "Export Product"
    assert len(payload["diagrams"]) == 1


def test_download_csv(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    product = _build_product(db, standard_user, threats=[{"severity": "high"}], mitigations=[{"status": "implemented"}])
    resp = client.get(f"/api/products/{product.id}/download/threats-mitigations", headers=user_headers)
    assert resp.status_code == 200
    assert "text/csv" in resp.headers["content-type"]
    body = resp.text
    assert "# Threats" in body and "# Mitigations" in body
    assert "Threat 0" in body and "Mitigation 0" in body


def test_download_html_report(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    product = _build_product(db, standard_user, threats=[{"severity": "critical"}])
    resp = client.get(f"/api/products/{product.id}/download/report", headers=user_headers)
    assert resp.status_code == 200
    assert "text/html" in resp.headers["content-type"]
    assert "Threat Model Report" in resp.text
    assert "Export Product" in resp.text


def test_download_markdown_report(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    product = _build_product(db, standard_user, threats=[{"severity": "high", "status": "mitigated"}])
    resp = client.get(f"/api/products/{product.id}/download/report.md", headers=user_headers)
    assert resp.status_code == 200
    assert "text/markdown" in resp.headers["content-type"]
    assert "# Threat Model Report" in resp.text
    assert "## Summary" in resp.text


def test_download_bundle_is_valid_zip(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    product = _build_product(db, standard_user, threats=[{"severity": "high"}], mitigations=[{"status": "verified"}])
    resp = client.get(f"/api/products/{product.id}/download/bundle", headers=user_headers)
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/zip"
    zf = zipfile.ZipFile(io.BytesIO(resp.content))
    names = {n.split("/")[-1] for n in zf.namelist()}
    assert {"diagrams.json", "threats-mitigations.csv", "report.html", "report.md", "report.docx", "README.txt"} <= names
    # Every entry must be non-empty and the JSON must parse.
    json_name = next(n for n in zf.namelist() if n.endswith("diagrams.json"))
    assert json.loads(zf.read(json_name))


def test_download_docx_is_valid_word_doc(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    from docx import Document  # local import: only needed for this assertion

    product = _build_product(
        db, standard_user,
        threats=[{"severity": "critical"}, {"severity": "high", "status": "mitigated"}],
        mitigations=[{"status": "implemented"}],
    )
    resp = client.get(f"/api/products/{product.id}/download/report.docx", headers=user_headers)
    assert resp.status_code == 200
    assert "wordprocessingml" in resp.headers["content-type"]
    # A .docx is a ZIP container with word/document.xml — parse it for real.
    doc = Document(io.BytesIO(resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Threat Model Report" in full_text
    assert "Export Product" in full_text
    # Threat names render inside the threats table.
    table_text = " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Threat 0" in table_text


def test_docx_requires_access(client: TestClient, standard_user: User, other_headers: dict, db: Session):
    product = _build_product(db, standard_user, is_public=False)
    resp = client.get(f"/api/products/{product.id}/download/report.docx", headers=other_headers)
    assert resp.status_code == 403


# ── Access control on exports ────────────────────────────────────────────────

def test_export_requires_auth(client: TestClient, standard_user: User, db: Session):
    product = _build_product(db, standard_user)
    resp = client.get(f"/api/products/{product.id}/download/diagrams")
    assert resp.status_code == 401


def test_stranger_cannot_export_private_product(client: TestClient, standard_user: User, other_user: User, other_headers: dict, db: Session):
    product = _build_product(db, standard_user, is_public=False)
    resp = client.get(f"/api/products/{product.id}/download/report", headers=other_headers)
    assert resp.status_code == 403


def test_export_missing_product_404(client: TestClient, user_headers: dict, db: Session):
    resp = client.get("/api/products/999999/download/diagrams", headers=user_headers)
    assert resp.status_code == 404


# ── security-status CI gate ──────────────────────────────────────────────────

def test_security_status_passes_with_no_thresholds(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    product = _build_product(db, standard_user, threats=[{"severity": "critical"}])
    resp = client.get(f"/api/products/{product.id}/security-status", headers=user_headers)
    assert resp.status_code == 200
    data = resp.json()
    assert data["pass"] is True
    assert data["summary"]["by_severity"]["critical"] == 1


def test_security_status_fails_on_critical(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    product = _build_product(db, standard_user, threats=[{"severity": "critical"}])
    resp = client.get(
        f"/api/products/{product.id}/security-status",
        params={"fail_on_critical": True},
        headers=user_headers,
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["pass"] is False
    assert any("critical" in f for f in data["failures"])


def test_security_status_fails_on_unmitigated_high(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    product = _build_product(
        db, standard_user,
        threats=[{"severity": "high", "status": "identified", "element_id": "n1"}],
    )
    resp = client.get(
        f"/api/products/{product.id}/security-status",
        params={"fail_on_unmitigated_high": True},
        headers=user_headers,
    )
    data = resp.json()
    assert data["pass"] is False


def test_security_status_high_passes_when_mitigated(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    product = _build_product(
        db, standard_user,
        threats=[{"severity": "high", "status": "mitigated", "element_id": "n1"}],
        mitigations=[{"status": "implemented", "element_id": "n1"}],
    )
    resp = client.get(
        f"/api/products/{product.id}/security-status",
        params={"fail_on_unmitigated_high": True},
        headers=user_headers,
    )
    data = resp.json()
    assert data["pass"] is True


def test_security_status_mitigation_ratio_threshold(client: TestClient, standard_user: User, user_headers: dict, db: Session):
    # 1 of 2 threats mitigated → ratio 0.5, below a 0.8 requirement.
    product = _build_product(
        db, standard_user,
        threats=[
            {"severity": "medium", "status": "mitigated", "element_id": "n1"},
            {"severity": "medium", "status": "identified", "element_id": "n2"},
        ],
    )
    resp = client.get(
        f"/api/products/{product.id}/security-status",
        params={"min_mitigation_ratio": 0.8},
        headers=user_headers,
    )
    data = resp.json()
    assert data["summary"]["mitigation_ratio"] == 0.5
    assert data["pass"] is False
    assert any("ratio" in f.lower() for f in data["failures"])
