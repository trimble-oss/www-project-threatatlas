"""Download endpoints for a product: diagrams JSON, threats/mitigations CSV,
standalone HTML report, a native Word (.docx) report, and a ZIP bundle
containing all of the above.

The HTML report is intentionally print-friendly so users can open it in a
browser and Save-as-PDF without needing server-side PDF tooling; the DOCX
report is generated server-side for audit deliverables that need a native
editable document.
"""

from __future__ import annotations

import csv
import html
import io
import json
import zipfile
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor

from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import Response
from sqlalchemy.orm import Session, joinedload

from app.database import get_db
from app.models import (
    Diagram,
    DiagramMitigation,
    DiagramThreat,
    Product as ProductModel,
    User as UserModel,
)
from app.auth.dependencies import get_current_user
from app.auth.permissions import PermissionDenied, can_access_product

router = APIRouter(prefix="/products", tags=["products"])


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _load_product(product_id: int, user: UserModel, db: Session) -> ProductModel:
    product = (
        db.query(ProductModel)
        .options(
            joinedload(ProductModel.diagrams).joinedload(Diagram.diagram_threats).joinedload(DiagramThreat.threat),
            joinedload(ProductModel.diagrams).joinedload(Diagram.diagram_mitigations).joinedload(DiagramMitigation.mitigation),
            joinedload(ProductModel.collaborators),
        )
        .filter(ProductModel.id == product_id)
        .first()
    )
    if not product:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Product not found")
    if not can_access_product(user, product):
        raise PermissionDenied("Not authorized to access this product")
    return product


def _safe_filename(name: str) -> str:
    return "".join(c if c.isalnum() or c in "-_" else "_" for c in name).strip("_") or "product"


def _diagrams_payload(product: ProductModel) -> dict:
    return {
        "product": {
            "id": product.id,
            "name": product.name,
            "description": product.description,
            "status": product.status,
            "business_area": product.business_area,
            "owner_name": product.owner_name,
            "owner_email": product.owner_email,
            "repository_url": product.repository_url,
            "application_url": product.application_url,
            "confluence_url": product.confluence_url,
        },
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "diagrams": [
            {
                "id": d.id,
                "name": d.name,
                "description": d.description,
                "current_version": d.current_version,
                "created_at": d.created_at.isoformat() if d.created_at else None,
                "data": d.diagram_data or {},
            }
            for d in product.diagrams
        ],
    }


def _threats_mitigations_rows(product: ProductModel) -> tuple[list[dict], list[dict]]:
    threats: list[dict] = []
    mitigations: list[dict] = []
    for d in product.diagrams:
        for dt in d.diagram_threats:
            threats.append({
                "diagram": d.name,
                "element_id": dt.element_id,
                "element_type": dt.element_type,
                "threat_name": dt.threat.name if dt.threat else "",
                "category": dt.threat.category if dt.threat else "",
                "description": dt.threat.description if dt.threat else "",
                "status": dt.status,
                "severity": dt.severity,
                "likelihood": dt.likelihood,
                "impact": dt.impact,
                "risk_score": dt.risk_score,
                "comments": dt.comments or "",
            })
        for dm in d.diagram_mitigations:
            mitigations.append({
                "diagram": d.name,
                "element_id": dm.element_id,
                "element_type": dm.element_type,
                "mitigation_name": dm.mitigation.name if dm.mitigation else "",
                "category": dm.mitigation.category if dm.mitigation else "",
                "description": dm.mitigation.description if dm.mitigation else "",
                "status": dm.status,
                "linked_threat_id": dm.threat_id,
                "comments": dm.comments or "",
            })
    return threats, mitigations


def _threats_csv(threats: list[dict], mitigations: list[dict]) -> str:
    buf = io.StringIO()
    buf.write("# Threats\n")
    if threats:
        w = csv.DictWriter(buf, fieldnames=list(threats[0].keys()))
        w.writeheader()
        w.writerows(threats)
    else:
        buf.write("(no threats)\n")
    buf.write("\n# Mitigations\n")
    if mitigations:
        w = csv.DictWriter(buf, fieldnames=list(mitigations[0].keys()))
        w.writeheader()
        w.writerows(mitigations)
    else:
        buf.write("(no mitigations)\n")
    return buf.getvalue()


def _html_report(product: ProductModel, threats: list[dict], mitigations: list[dict]) -> str:
    def esc(v):
        return html.escape(str(v)) if v not in (None, "") else "—"

    metadata_rows = [
        ("Status", product.status),
        ("Business area", product.business_area),
        ("Owner", f"{product.owner_name or ''} {('<' + product.owner_email + '>') if product.owner_email else ''}".strip()),
        ("Repository", product.repository_url),
        ("Application URL", product.application_url),
        ("Confluence", product.confluence_url),
    ]
    meta_html = "".join(
        f"<tr><th>{esc(k)}</th><td>{esc(v)}</td></tr>" for k, v in metadata_rows if v
    )

    def threat_row(t):
        return (
            f"<tr>"
            f"<td>{esc(t['diagram'])}</td>"
            f"<td>{esc(t['element_id'])}</td>"
            f"<td>{esc(t['threat_name'])}</td>"
            f"<td>{esc(t['category'])}</td>"
            f"<td>{esc(t['status'])}</td>"
            f"<td>{esc(t['severity'])}</td>"
            f"<td>{esc(t['risk_score'])}</td>"
            f"<td>{esc(t['description'])}</td>"
            f"</tr>"
        )

    def mit_row(m):
        return (
            f"<tr>"
            f"<td>{esc(m['diagram'])}</td>"
            f"<td>{esc(m['element_id'])}</td>"
            f"<td>{esc(m['mitigation_name'])}</td>"
            f"<td>{esc(m['category'])}</td>"
            f"<td>{esc(m['status'])}</td>"
            f"<td>{esc(m['description'])}</td>"
            f"</tr>"
        )

    threats_html = "".join(threat_row(t) for t in threats) or '<tr><td colspan="8">No threats recorded.</td></tr>'
    mitigations_html = "".join(mit_row(m) for m in mitigations) or '<tr><td colspan="6">No mitigations recorded.</td></tr>'

    diagrams_html = "".join(
        f"<li><strong>{esc(d.name)}</strong> — {esc(d.description) if d.description else 'no description'} "
        f"<em>(v{d.current_version})</em></li>"
        for d in product.diagrams
    ) or "<li>No diagrams.</li>"

    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Threat Model Report — {esc(product.name)}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
          margin: 40px auto; max-width: 1100px; color: #222; }}
  h1 {{ border-bottom: 3px solid #3b82f6; padding-bottom: 10px; }}
  h2 {{ margin-top: 40px; border-bottom: 1px solid #e5e7eb; padding-bottom: 6px; }}
  table {{ border-collapse: collapse; width: 100%; margin-top: 12px; font-size: 13px; }}
  th, td {{ border: 1px solid #d1d5db; padding: 8px 10px; text-align: left; vertical-align: top; }}
  th {{ background: #f3f4f6; font-weight: 600; }}
  .meta th {{ width: 160px; background: #f9fafb; }}
  .footer {{ margin-top: 60px; color: #6b7280; font-size: 11px; text-align: center; }}
  @media print {{ body {{ margin: 20px; }} h1 {{ page-break-after: avoid; }} table {{ page-break-inside: avoid; }} }}
</style>
</head>
<body>
<h1>Threat Model Report — {esc(product.name)}</h1>
<p><em>{esc(product.description or "")}</em></p>

<h2>Project Metadata</h2>
<table class="meta">{meta_html or '<tr><td>No project metadata provided.</td></tr>'}</table>

<h2>Diagrams</h2>
<ul>{diagrams_html}</ul>

<h2>Threats ({len(threats)})</h2>
<table>
  <thead><tr>
    <th>Diagram</th><th>Element</th><th>Threat</th><th>Category</th>
    <th>Status</th><th>Severity</th><th>Risk</th><th>Description</th>
  </tr></thead>
  <tbody>{threats_html}</tbody>
</table>

<h2>Mitigations ({len(mitigations)})</h2>
<table>
  <thead><tr>
    <th>Diagram</th><th>Element</th><th>Mitigation</th><th>Category</th>
    <th>Status</th><th>Description</th>
  </tr></thead>
  <tbody>{mitigations_html}</tbody>
</table>

<p class="footer">Generated by ThreatAtlas on {generated}. Use your browser's Print → Save as PDF to export.</p>
</body>
</html>
"""


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@router.get("/{product_id}/download/diagrams")
def download_diagrams(product_id: int, current_user: UserModel = Depends(get_current_user), db: Session = Depends(get_db)):
    product = _load_product(product_id, current_user, db)
    body = json.dumps(_diagrams_payload(product), indent=2, default=str)
    return Response(
        content=body,
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(product.name)}-diagrams.json"'},
    )


@router.get("/{product_id}/download/threats-mitigations")
def download_threats_mitigations(product_id: int, current_user: UserModel = Depends(get_current_user), db: Session = Depends(get_db)):
    product = _load_product(product_id, current_user, db)
    threats, mitigations = _threats_mitigations_rows(product)
    return Response(
        content=_threats_csv(threats, mitigations),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(product.name)}-threats-mitigations.csv"'},
    )


@router.get("/{product_id}/download/report")
def download_report(product_id: int, current_user: UserModel = Depends(get_current_user), db: Session = Depends(get_db)):
    product = _load_product(product_id, current_user, db)
    threats, mitigations = _threats_mitigations_rows(product)
    return Response(
        content=_html_report(product, threats, mitigations),
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(product.name)}-report.html"'},
    )


def _markdown_report(product: ProductModel, threats: list[dict], mitigations: list[dict]) -> str:
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    lines: list[str] = [
        f"# Threat Model Report — {product.name}",
        f"",
        f"**Generated:** {now}  ",
    ]
    if product.status:
        lines.append(f"**Status:** {product.status}  ")
    if product.business_area:
        lines.append(f"**Business Area:** {product.business_area}  ")
    if product.owner_name:
        lines.append(f"**Owner:** {product.owner_name}" + (f" <{product.owner_email}>" if product.owner_email else "") + "  ")
    if product.repository_url:
        lines.append(f"**Repository:** {product.repository_url}  ")
    if product.confluence_url:
        lines.append(f"**Confluence:** {product.confluence_url}  ")
    lines.append("")

    # Summary stats
    total_t = len(threats)
    critical = sum(1 for t in threats if t.get("severity") == "critical")
    high = sum(1 for t in threats if t.get("severity") == "high")
    medium = sum(1 for t in threats if t.get("severity") == "medium")
    low = sum(1 for t in threats if t.get("severity") == "low")
    mitigated = sum(1 for t in threats if t.get("status") == "mitigated")
    active_mits = sum(1 for m in mitigations if m.get("status") in ("implemented", "verified"))
    ratio = round((mitigated / total_t * 100)) if total_t else 0

    lines += [
        "## Summary",
        "",
        f"| Metric | Value |",
        f"|--------|-------|",
        f"| Total threats | {total_t} |",
        f"| Critical | {critical} |",
        f"| High | {high} |",
        f"| Medium | {medium} |",
        f"| Low | {low} |",
        f"| Mitigated threats | {mitigated} ({ratio}%) |",
        f"| Active mitigations | {active_mits} / {len(mitigations)} |",
        "",
    ]

    # Threats by diagram
    diagrams_seen: dict[str, list[dict]] = {}
    for t in threats:
        diagrams_seen.setdefault(t["diagram"], []).append(t)

    lines.append("## Threats")
    lines.append("")
    for diag_name, diag_threats in diagrams_seen.items():
        lines.append(f"### {diag_name}")
        lines.append("")
        lines.append("| Element | Threat | Category | Severity | Status | Risk Score |")
        lines.append("|---------|--------|----------|----------|--------|------------|")
        for t in diag_threats:
            sev = t.get("severity") or "—"
            rs = str(t.get("risk_score")) if t.get("risk_score") is not None else "—"
            lines.append(
                f"| {t['element_id']} | {t['threat_name']} | {t['category']} "
                f"| {sev} | {t['status']} | {rs} |"
            )
        lines.append("")

    # Mitigations
    lines.append("## Mitigations")
    lines.append("")
    if mitigations:
        mit_by_diag: dict[str, list[dict]] = {}
        for m in mitigations:
            mit_by_diag.setdefault(m["diagram"], []).append(m)
        for diag_name, diag_mits in mit_by_diag.items():
            lines.append(f"### {diag_name}")
            lines.append("")
            lines.append("| Element | Mitigation | Category | Status |")
            lines.append("|---------|-----------|----------|--------|")
            for m in diag_mits:
                lines.append(f"| {m['element_id']} | {m['mitigation_name']} | {m['category']} | {m['status']} |")
            lines.append("")
    else:
        lines.append("_No mitigations recorded._\n")

    lines.append("---")
    lines.append(f"_Exported from ThreatAtlas_")
    return "\n".join(lines)


_SEVERITY_COLORS = {
    "critical": RGBColor(0xB9, 0x1C, 0x1C),
    "high": RGBColor(0xC2, 0x41, 0x0C),
    "medium": RGBColor(0xB4, 0x53, 0x09),
    "low": RGBColor(0x37, 0x7F, 0x3B),
}


def _docx_report(product: ProductModel, threats: list[dict], mitigations: list[dict]) -> bytes:
    """Render a native Word (.docx) threat-model report and return its bytes."""
    doc = Document()
    doc.add_heading(f"Threat Model Report — {product.name}", level=0)
    if product.description:
        doc.add_paragraph(product.description).italic = True

    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta_p = doc.add_paragraph()
    meta_p.add_run("Generated: ").bold = True
    meta_p.add_run(generated)

    # Project metadata
    metadata_rows = [
        ("Status", product.status),
        ("Business area", product.business_area),
        ("Owner", product.owner_name),
        ("Owner email", product.owner_email),
        ("Repository", product.repository_url),
        ("Application URL", product.application_url),
        ("Confluence", product.confluence_url),
    ]
    metadata_rows = [(k, v) for k, v in metadata_rows if v]
    if metadata_rows:
        doc.add_heading("Project Metadata", level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for key, value in metadata_rows:
            cells = table.add_row().cells
            cells[0].text = key
            cells[1].text = str(value)

    # Summary
    total = len(threats)
    by_sev = {s: sum(1 for t in threats if t.get("severity") == s) for s in ("critical", "high", "medium", "low")}
    mitigated = sum(1 for t in threats if t.get("status") == "mitigated")
    doc.add_heading("Summary", level=1)
    summary = doc.add_table(rows=0, cols=2)
    summary.style = "Light Grid Accent 1"
    for label, value in [
        ("Total threats", total),
        ("Critical", by_sev["critical"]),
        ("High", by_sev["high"]),
        ("Medium", by_sev["medium"]),
        ("Low", by_sev["low"]),
        ("Mitigated threats", f"{mitigated} ({round(mitigated / total * 100) if total else 0}%)"),
        ("Total mitigations", len(mitigations)),
    ]:
        cells = summary.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    # Threats
    doc.add_heading(f"Threats ({total})", level=1)
    if threats:
        t_table = doc.add_table(rows=1, cols=6)
        t_table.style = "Light Grid Accent 1"
        for i, header in enumerate(["Diagram", "Element", "Threat", "Category", "Severity", "Status"]):
            run = t_table.rows[0].cells[i].paragraphs[0].add_run(header)
            run.bold = True
        for t in threats:
            cells = t_table.add_row().cells
            cells[0].text = str(t.get("diagram", ""))
            cells[1].text = str(t.get("element_id", ""))
            cells[2].text = str(t.get("threat_name", ""))
            cells[3].text = str(t.get("category", ""))
            sev = t.get("severity") or "—"
            sev_run = cells[4].paragraphs[0].add_run(sev)
            if sev in _SEVERITY_COLORS:
                sev_run.font.color.rgb = _SEVERITY_COLORS[sev]
                sev_run.bold = True
            cells[5].text = str(t.get("status", ""))
    else:
        doc.add_paragraph("No threats recorded.")

    # Mitigations
    doc.add_heading(f"Mitigations ({len(mitigations)})", level=1)
    if mitigations:
        m_table = doc.add_table(rows=1, cols=5)
        m_table.style = "Light Grid Accent 1"
        for i, header in enumerate(["Diagram", "Element", "Mitigation", "Category", "Status"]):
            run = m_table.rows[0].cells[i].paragraphs[0].add_run(header)
            run.bold = True
        for m in mitigations:
            cells = m_table.add_row().cells
            cells[0].text = str(m.get("diagram", ""))
            cells[1].text = str(m.get("element_id", ""))
            cells[2].text = str(m.get("mitigation_name", ""))
            cells[3].text = str(m.get("category", ""))
            cells[4].text = str(m.get("status", ""))
    else:
        doc.add_paragraph("No mitigations recorded.")

    footer = doc.add_paragraph()
    footer_run = footer.add_run(f"Generated by ThreatAtlas on {generated}.")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _security_status_payload(
    product: ProductModel,
    threats: list[dict],
    mitigations: list[dict],
    fail_on_critical: bool,
    fail_on_unmitigated_high: bool,
    min_mitigation_ratio: float | None,
) -> dict:
    total = len(threats)
    by_sev = {
        "critical": sum(1 for t in threats if t.get("severity") == "critical"),
        "high": sum(1 for t in threats if t.get("severity") == "high"),
        "medium": sum(1 for t in threats if t.get("severity") == "medium"),
        "low": sum(1 for t in threats if t.get("severity") == "low"),
        "unscored": sum(1 for t in threats if not t.get("severity")),
    }
    mitigated = sum(1 for t in threats if t.get("status") == "mitigated")
    active_mits = sum(1 for m in mitigations if m.get("status") in ("implemented", "verified"))
    ratio = round(mitigated / total, 4) if total else 1.0

    # Unmitigated high-severity threats
    mitigated_element_ids = {m["element_id"] for m in mitigations if m.get("status") in ("implemented", "verified")}
    unmitigated_high = [
        t for t in threats
        if t.get("severity") in ("high", "critical") and t.get("status") != "mitigated"
        and t["element_id"] not in mitigated_element_ids
    ]

    failures: list[str] = []
    if fail_on_critical and by_sev["critical"] > 0:
        failures.append(f"{by_sev['critical']} critical threat(s) require immediate attention")
    if fail_on_unmitigated_high and unmitigated_high:
        failures.append(f"{len(unmitigated_high)} high/critical threat(s) have no active mitigation")
    if min_mitigation_ratio is not None and ratio < min_mitigation_ratio:
        failures.append(
            f"Mitigation ratio {ratio:.0%} is below required threshold {min_mitigation_ratio:.0%}"
        )

    return {
        "product_id": product.id,
        "product_name": product.name,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "thresholds": {
            "fail_on_critical": fail_on_critical,
            "fail_on_unmitigated_high": fail_on_unmitigated_high,
            "min_mitigation_ratio": min_mitigation_ratio,
        },
        "summary": {
            "total_threats": total,
            "by_severity": by_sev,
            "total_mitigations": len(mitigations),
            "active_mitigations": active_mits,
            "mitigated_threats": mitigated,
            "mitigation_ratio": ratio,
        },
        "pass": len(failures) == 0,
        "failures": failures,
    }


@router.get("/{product_id}/download/bundle")
def download_bundle(product_id: int, current_user: UserModel = Depends(get_current_user), db: Session = Depends(get_db)):
    product = _load_product(product_id, current_user, db)
    threats, mitigations = _threats_mitigations_rows(product)
    slug = _safe_filename(product.name)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{slug}/diagrams.json", json.dumps(_diagrams_payload(product), indent=2, default=str))
        zf.writestr(f"{slug}/threats-mitigations.csv", _threats_csv(threats, mitigations))
        zf.writestr(f"{slug}/report.html", _html_report(product, threats, mitigations))
        zf.writestr(f"{slug}/report.md", _markdown_report(product, threats, mitigations))
        zf.writestr(f"{slug}/report.docx", _docx_report(product, threats, mitigations))
        zf.writestr(
            f"{slug}/README.txt",
            f"ThreatAtlas export for product: {product.name}\n"
            f"Generated: {datetime.now(timezone.utc).isoformat()}\n\n"
            f"- diagrams.json: full diagram data (ReactFlow nodes/edges) for re-import\n"
            f"- threats-mitigations.csv: all threats and mitigations across every diagram\n"
            f"- report.html: standalone HTML report — open in browser and Print → Save as PDF\n"
            f"- report.md: Markdown report — paste into GitHub PRs, wikis, or CI summaries\n"
            f"- report.docx: native Word report for audit deliverables\n",
        )
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{slug}-threatatlas.zip"'},
    )


@router.get("/{product_id}/download/report.md")
def download_report_markdown(product_id: int, current_user: UserModel = Depends(get_current_user), db: Session = Depends(get_db)):
    product = _load_product(product_id, current_user, db)
    threats, mitigations = _threats_mitigations_rows(product)
    return Response(
        content=_markdown_report(product, threats, mitigations),
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(product.name)}-report.md"'},
    )


@router.get("/{product_id}/download/report.docx")
def download_report_docx(product_id: int, current_user: UserModel = Depends(get_current_user), db: Session = Depends(get_db)):
    product = _load_product(product_id, current_user, db)
    threats, mitigations = _threats_mitigations_rows(product)
    return Response(
        content=_docx_report(product, threats, mitigations),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(product.name)}-report.docx"'},
    )


@router.get("/{product_id}/security-status")
def security_status(
    product_id: int,
    fail_on_critical: bool = Query(default=False, description="Fail if any critical threats exist"),
    fail_on_unmitigated_high: bool = Query(default=False, description="Fail if high/critical threats have no active mitigation"),
    min_mitigation_ratio: float | None = Query(default=None, description="Minimum required mitigation ratio (0.0–1.0)"),
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Machine-readable security posture check for CI/CD pipelines.

    Returns HTTP 200 with pass=true when all thresholds are met,
    or HTTP 200 with pass=false when any threshold is violated.
    Consumers should check the `pass` field — not the HTTP status code.
    """
    product = _load_product(product_id, current_user, db)
    threats, mitigations = _threats_mitigations_rows(product)
    payload = _security_status_payload(
        product, threats, mitigations,
        fail_on_critical, fail_on_unmitigated_high, min_mitigation_ratio,
    )
    return payload
